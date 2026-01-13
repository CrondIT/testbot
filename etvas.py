import openai
from telegram import Update
from telegram.ext import Updater, CommandHandler, MessageHandler, Filters, CallbackContext
from telegram import InputFile
from docx import Document

# Замените на ваш API-ключ OpenAI
openai.api_key = "your_openai_api_key"

# Функция для общения с GPT-5.2
def generate_text_from_gpt(prompt: str) -> str:
    response = openai.Completion.create(
        engine="gpt-5.2",  # Убедитесь, что используете правильную модель
        prompt=prompt,
        max_tokens=1500,
        temperature=0.7
    )
    return response.choices[0].text.strip()

# Функция для создания DOCX файла
def create_docx(content: str, filename: str) -> str:
    document = Document()
    document.add_heading('Ответ от GPT', 0)
    document.add_paragraph(content)
    
    # Сохраняем файл на диск
    document.save(filename)
    return filename

# Функция обработки сообщения с запросом на docx
def handle_message(update: Update, context: CallbackContext) -> None:
    user_message = update.message.text
    
    if "формат docx" in user_message.lower():
        # Отправляем запрос GPT-5.2 для обработки текста
        response_text = generate_text_from_gpt(user_message)
        
        # Создаём DOCX файл
        filename = "output.docx"
        create_docx(response_text, filename)
        
        # Отправляем файл пользователю
        with open(filename, "rb") as file:
            update.message.reply_document(document=file)

# Функция обработки команды start
def start(update: Update, context: CallbackContext) -> None:
    update.message.reply_text("Привет! Напиши мне текст, и если тебе нужно, я отправлю ответ в формате DOCX.")

def main():
    # Токен вашего бота
    token = 'your_telegram_bot_token'
    
    # Инициализация бота
    updater = Updater(token, use_context=True)
    dispatcher = updater.dispatcher
    
    # Обработчики команд
    dispatcher.add_handler(CommandHandler("start", start))
    
    # Обработчик сообщений
    dispatcher.add_handler(MessageHandler(Filters.text & ~Filters.command, handle_message))
    
    # Запуск бота
    updater.start_polling()
    updater.idle()

if __name__ == '__main__':
    main()
