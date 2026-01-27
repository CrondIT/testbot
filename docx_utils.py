"""Utility functions for creating and handling DOCX files."""

import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Mm
from docx.oxml.shared import OxmlElement, qn
import io
import json
from telegram import InputFile
from telegram.helpers import escape_markdown
from telegram.error import TimedOut
import matplotlib
import matplotlib.pyplot as plt
import numpy as np
from global_state import DOCUMENT_JSON_SCHEMA

# Константы для формата страницы и полей (в миллиметрах)
DEFAULT_PAGE_WIDTH_MM = 210  # Ширина A4 в мм
DEFAULT_PAGE_HEIGHT_MM = 297  # Высота A4 в мм
DEFAULT_LEFT_MARGIN_MM = 20  # Левое поле по умолчанию
DEFAULT_RIGHT_MARGIN_MM = 10  # Правое поле по умолчанию
DEFAULT_TOP_MARGIN_MM = 20  # Верхнее поле по умолчанию
DEFAULT_BOTTOM_MARGIN_MM = 20  # Нижнее поле по умолчанию

# Константы для размеров колонтитулов (в миллиметрах)
DEFAULT_HEADER_HEIGHT_MM = 10  # Высота верхнего колонтитула по умолчанию
DEFAULT_FOOTER_HEIGHT_MM = 10  # Высота нижнего колонтитула по умолчанию


def clean_html_tags(text):
    """
    Очищает текст от HTML-тегов и других форматирований,
    не поддерживаемых в python-docx.

    Args:
        text (str): Текст для очистки

    Returns:
        str: Очищенный текст
    """
    if not isinstance(text, str):
        return str(text)

    # Удаляем span теги и их атрибуты
    text = re.sub(r"<span[^>]*>", "", text)
    text = re.sub(r"</span>", "", text)

    # Удаляем другие теги
    text = re.sub(r"<[^>]+>", "", text)

    # Удаляем двойные звездочки (жирный текст в markdown)
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)

    # Удаляем одинарные звездочки (курсив в markdown)
    text = re.sub(r"(?<!\*)\*([^\*]+?)\*(?!\*)", r"\1", text)

    # Удаляем подчеркивания (курсив в markdown)
    text = re.sub(r"(?<!_)_([^_]+?)_(?!_)", r"\1", text)

    return text.strip()


matplotlib.use("Agg")  # Use non-interactive backend

# JSON_SCHEMA теперь определена в global_state.py
# и импортируется как DOCUMENT_JSON_SCHEMA
# Для обратной совместимости, создаем алиас
JSON_SCHEMA = DOCUMENT_JSON_SCHEMA


class DocxRenderer:
    def __init__(self):
        self.doc = Document()
        # Установка полей страницы по умолчанию
        section = self.doc.sections[0]
        section.page_width = Mm(DEFAULT_PAGE_WIDTH_MM)
        section.page_height = Mm(DEFAULT_PAGE_HEIGHT_MM)
        section.left_margin = Mm(DEFAULT_LEFT_MARGIN_MM)
        section.right_margin = Mm(DEFAULT_RIGHT_MARGIN_MM)
        section.top_margin = Mm(DEFAULT_TOP_MARGIN_MM)
        section.bottom_margin = Mm(DEFAULT_BOTTOM_MARGIN_MM)

    def render(self, data: dict, output):
        """
        Рендерим JSON документ в DOCX.
        :param data: dict с ключами "meta", "header", "footer" и "blocks"
        :param output: path или BytesIO
        """
        # Проверяем, есть ли в блоках заголовки (heading blocks)
        blocks = data.get("blocks", [])
        has_heading_blocks = any(
            block.get("type") == "heading" for block in blocks
        )

        # Отображаем мета-заголовок только если нет заголовков в блоках
        # (пользовательское оформление имеет приоритет)
        meta = data.get("meta", {})
        if not has_heading_blocks:
            self._render_meta(meta)
        else:
            # Если есть заголовки в блоках, не отображаем заголовок из мета
            # (пользовательское оформление имеет приоритет)
            pass

        # Обрабатываем верхний и нижний колонтитулы
        self._render_header_footer(data)

        for block in blocks:
            self._render_block(block)

        # Сохраняем в файл или BytesIO
        self.doc.save(output)

    def _render_meta(self, meta: dict):
        # Заголовок из метаданных отображается на странице по умолчанию,
        # если пользователь не указал, что не хочет его отображать
        # и если в блоках нет заголовков
        #  (пользовательское оформление имеет приоритет)
        title = meta.get("title")
        hide_title = meta.get("hide_title", False)

        # Мы не добавляем заголовок из мета, если в процессе рендеринга блоков
        # будет обнаружен заголовок
        # (пользовательское оформление имеет приоритет)
        # Но для совместимости, добавляем заголовок из мета, если он не скрыт
        # и если пользователь не добавил заголовок в блоки
        if title is not None and title != "" and not hide_title:
            # Note: We still add the meta title,
            # but users can override by using heading blocks
            title_heading = self.doc.add_heading(title, level=1)
            title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _render_header_footer(self, data: dict):
        """
        Рендерим верхний и нижний колонтитулы из JSON.
        :param data: dict с ключами "header" и "footer"
        """
        header_data = data.get("header", {})
        footer_data = data.get("footer", {})

        # Получаем секцию документа
        section = self.doc.sections[0]

        # Устанавливаем верхний колонтитул
        if header_data:
            header = section.header
            self._apply_header_footer_format(header, header_data, "header")

        # Устанавливаем нижний колонтитул
        if footer_data:
            footer = section.footer
            self._apply_header_footer_format(footer, footer_data, "footer")

    def _get_available_page_width(self):
        """
        Возвращает доступную ширину страницы
        (ширина страницы минус левые и правые поля).
        """
        section = self.doc.sections[0]
        page_width = section.page_width
        left_margin = section.left_margin
        right_margin = section.right_margin
        available_width = page_width - left_margin - right_margin
        return available_width

    def _apply_header_footer_format(
        self, header_footer_part, format_data: dict, part_type: str
    ):
        """
        Применяем форматирование к верхнему или нижнему колонтитулу.
        :param header_footer_part: объект header или footer
        :param format_data: словарь с параметрами форматирования
        :param part_type: тип части ('header' или 'footer')
        """
        content = format_data.get("content", "")
        font_name = format_data.get("font_name", None)
        font_size = format_data.get("font_size", None)
        color = format_data.get("color", None)
        bold = format_data.get("bold", None)
        italic = format_data.get("italic", None)
        alignment = format_data.get("alignment", "left")  # left, center, right

        # Очищаем существующий контент
        for paragraph in header_footer_part.paragraphs:
            p_element = paragraph._p
            p_element.getparent().remove(p_element)

        # Добавляем новый параграф с содержимым
        p = (
            header_footer_part.paragraphs[0]
            if header_footer_part.paragraphs
            else header_footer_part.add_paragraph()
        )

        # Проверяем, есть ли настройки номеров страниц для футера
        if part_type == "footer":
            page_number_settings = format_data.get("page_number", {})
            if page_number_settings.get("enabled", False):
                # Обработка номеров страниц
                self._add_page_numbers_to_footer(
                    p, content, page_number_settings
                )
            else:
                # Обычная обработка без номеров страниц
                run = p.add_run(content)
                self._apply_run_formatting(
                    run, font_name, font_size, bold, italic, color
                )
        else:
            # Обработка хедера (без номеров страниц)
            run = p.add_run(content)
            self._apply_run_formatting(
                run, font_name, font_size, bold, italic, color
            )

        # Применяем выравнивание
        if alignment == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif alignment == "justify":
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # left - по умолчанию

    def _apply_run_formatting(
        self, run, font_name, font_size, bold, italic, color
    ):
        """
        Применяет форматирование к run.
        """
        if font_name:
            run.font.name = font_name
        if font_size:
            from docx.shared import Pt

            run.font.size = Pt(font_size)
        if bold is not None:
            run.font.bold = bold
        if italic is not None:
            run.font.italic = italic
        if color:
            from docx.shared import RGBColor

            # Предполагаем, что цвет в формате "RRGGBB" или "RGB"
            if len(color) == 6:
                r = int(color[0:2], 16)
                g = int(color[2:4], 16)
                b = int(color[4:6], 16)
                run.font.color.rgb = RGBColor(r, g, b)

    def _add_page_numbers_to_footer(
        self, paragraph, content, page_number_settings
    ):
        """
        Добавляет номера страниц в футер.
        """
        from docx.oxml.shared import OxmlElement, qn

        # Получаем формат строки номера страницы
        format_str = page_number_settings.get(
            "format", "Стр. {PAGE} из {NUMPAGES}"
        )

        # Применяем форматирование к первому run в параграфе
        if paragraph.runs:
            first_run = paragraph.runs[0]
            font_name = page_number_settings.get("font_name")
            font_size = page_number_settings.get("font_size")
            bold = page_number_settings.get("bold")
            italic = page_number_settings.get("italic")
            color = page_number_settings.get("color")

            self._apply_run_formatting(
                first_run, font_name, font_size, bold, italic, color
            )

        # Заменяем плейсхолдеры на поля Word
        # Сначала заменим {PAGE} на поле номера страницы
        if "{PAGE}" in format_str:
            # Разбиваем строку на части до и после {PAGE}
            parts = format_str.split("{PAGE}")
            for i, part in enumerate(parts):
                if part:
                    paragraph.add_run(part)

                # Добавляем поле номера страницы, если это не последняя часть
                if i < len(parts) - 1:
                    # Добавляем поле номера текущей страницы
                    pg_num_run = paragraph.add_run()
                    fldChar = OxmlElement("w:fldChar")
                    fldChar.set(qn("w:fldCharType"), "begin")
                    pg_num_run._r.append(fldChar)

                    instrText = OxmlElement("w:instrText")
                    instrText.set(qn("xml:space"), "preserve")
                    instrText.text = "PAGE"
                    pg_num_run._r.append(instrText)

                    fldChar2 = OxmlElement("w:fldChar")
                    fldChar2.set(qn("w:fldCharType"), "end")
                    pg_num_run._r.append(fldChar2)

        # Если в строке есть {NUMPAGES}, добавляем поле общего числа страниц
        if "{NUMPAGES}" in format_str:
            # В этом случае нужно более сложное форматирование
            # Пересоздадим параграф с правильной структурой
            paragraph.clear()  # Очищаем текущий параграф

            # Разбиваем формат по плейсхолдерам
            import re

            # Ищем все плейсхолдеры
            tokens = re.split(r"(\{PAGE\}|\{NUMPAGES\})", format_str)

            for token in tokens:
                if token == "{PAGE}":
                    # Добавляем поле номера страницы
                    pg_num_run = paragraph.add_run()
                    fldChar = OxmlElement("w:fldChar")
                    fldChar.set(qn("w:fldCharType"), "begin")
                    pg_num_run._r.append(fldChar)

                    instrText = OxmlElement("w:instrText")
                    instrText.set(qn("xml:space"), "preserve")
                    instrText.text = "PAGE"
                    pg_num_run._r.append(instrText)

                    fldChar2 = OxmlElement("w:fldChar")
                    fldChar2.set(qn("w:fldCharType"), "end")
                    pg_num_run._r.append(fldChar2)

                elif token == "{NUMPAGES}":
                    # Добавляем поле общего числа страниц
                    pg_tot_run = paragraph.add_run()
                    fldChar = OxmlElement("w:fldChar")
                    fldChar.set(qn("w:fldCharType"), "begin")
                    pg_tot_run._r.append(fldChar)

                    instrText = OxmlElement("w:instrText")
                    instrText.set(qn("xml:space"), "preserve")
                    instrText.text = "NUMPAGES"
                    pg_tot_run._r.append(instrText)

                    fldChar2 = OxmlElement("w:fldChar")
                    fldChar2.set(qn("w:fldCharType"), "end")
                    pg_tot_run._r.append(fldChar2)

                else:
                    # Обычный текст
                    if token:
                        paragraph.add_run(token)

        # Применяем форматирование к первому run в параграфе
        if paragraph.runs:
            first_run = paragraph.runs[0]
            font_name = page_number_settings.get("font_name")
            font_size = page_number_settings.get("font_size")
            bold = page_number_settings.get("bold")
            italic = page_number_settings.get("italic")
            color = page_number_settings.get("color")

            self._apply_run_formatting(
                first_run, font_name, font_size, bold, italic, color
            )

    def _render_block(self, block: dict):
        block_type = block.get("type")

        if block_type == "heading":
            self._heading(block)
        elif block_type == "paragraph":
            self._paragraph(block)
        elif block_type == "list":
            self._list(block)
        elif block_type == "table":
            self._table(block)
        elif block_type == "math":
            self._math(block)
        elif block_type == "function_graph":
            self._function_graph(block)
        elif block_type == "toc":
            self._toc(block)
        else:
            raise ValueError(f"Unknown block type: {block_type}")

    def _toc(self, block: dict):
        """
        Создает оглавление документа.
        JSON-схема:
        {
            "type": "toc",
            "title": "Оглавление",  // Заголовок оглавления
            "levels": [1, 2, 3],    // Уровни заголовков для включения
            "font_name": "string",  // Шрифт для оглавления
            "font_size": 12,       // Размер шрифта
            "indent": 10,          // Отступ для вложенных уровней
            "leader_dots": true,   // Пунктирная линия до номера страницы
            "include_pages": true  // Включать номера страниц
        }
        """
        title = block.get("title", "Оглавление")
        levels = block.get("levels", [1, 2, 3])
        font_name = block.get("font_name", "Arial")
        font_size = block.get("font_size", 12)
        indent = block.get("indent", 10)
        leader_dots = block.get("leader_dots", True)
        include_pages = block.get("include_pages", True)

        # Добавляем заголовок оглавления
        toc_title = self.doc.add_heading(title, level=1)
        if font_name:
            if toc_title.runs:
                toc_title.runs[0].font.name = font_name
            else:
                toc_title.add_run().font.name = font_name
        if font_size:
            from docx.shared import Pt

            if toc_title.runs:
                toc_title.runs[0].font.size = Pt(font_size)

        # В текущей реализации python-docx не позволяет легко создавать
        # настоящие оглавления с гиперссылками, но мы можем создать
        # структурированный список, имитирующий оглавление

        # Вместо фиксированных записей, позволим
        # пользователю передать записи оглавления
        # через параметр entries, если он есть, иначе используем примеры
        toc_entries = block.get(
            "entries",
            [
                {"text": "Введение", "level": 1, "page": 1},
                {"text": "Основная часть", "level": 1, "page": 3},
                {"text": "Подраздел 1", "level": 2, "page": 4},
                {"text": "Подраздел 2", "level": 2, "page": 6},
                {"text": "Заключение", "level": 1, "page": 8},
                {"text": "Список литературы", "level": 1, "page": 10},
            ],
        )

        for entry in toc_entries:
            level = entry["level"]
            text = entry["text"]
            page = entry.get("page", 1)  # по умолчанию страница 1

            # Проверяем, входит ли уровень в разрешенные
            if level in levels:
                p = self.doc.add_paragraph()

                # Добавляем отступ в зависимости от уровня
                indent_points = (level - 1) * indent
                from docx.shared import Inches

                p.paragraph_format.left_indent = Inches(indent_points / 10.0)

                # Добавляем текст заголовка
                run = p.add_run(text)
                if font_name:
                    run.font.name = font_name
                if font_size:
                    from docx.shared import Pt

                    run.font.size = Pt(font_size)

                # Добавляем точки и номер страницы
                if include_pages and leader_dots:
                    # Вычисляем, сколько точек нужно добавить
                    # в зависимости от длины текста
                    # и ширины страницы
                    dots_needed = max(0, 60 - len(text) - len(str(page)))
                    dots = "." * min(
                        dots_needed, 30
                    )  # Ограничиваем количество точек

                    p.add_run(f" {dots} ")
                    p.add_run(str(page))

    def _heading(self, block: dict):
        level = block.get("level", 1)
        text = block.get("text", "")
        # Используем параметры из JSON, если они есть, иначе - стандартные
        font_name = block.get("font_name", None)
        font_size = block.get("font_size", None)
        bold = block.get("bold", None)
        italic = block.get("italic", None)
        color = block.get("color", None)

        heading = self.doc.add_heading(text, level=level)

        # Применяем стили, если они указаны в JSON
        if font_name or font_size or bold is not None or italic is not None:
            run = heading.runs[0] if heading.runs else heading.add_run(text)
            if font_name:
                run.font.name = font_name
            if font_size:
                from docx.shared import Pt

                run.font.size = Pt(font_size)
            if bold is not None:
                run.font.bold = bold
            if italic is not None:
                run.font.italic = italic
            if color:
                # Для цвета нужно импортировать RGBColor
                from docx.shared import RGBColor

                # Предполагаем, что цвет в формате "RRGGBB" или "RGB"
                if len(color) == 6:
                    r = int(color[0:2], 16)
                    g = int(color[2:4], 16)
                    b = int(color[4:6], 16)
                    run.font.color.rgb = RGBColor(r, g, b)

    def _paragraph(self, block: dict):
        # Используем параметры из JSON, если они есть, иначе - стандартные
        text = clean_html_tags(block.get("text", ""))
        font_name = block.get("font_name", None)
        font_size = block.get("font_size", None)
        left_indent = block.get("left_indent", None)
        right_indent = block.get("right_indent", None)
        space_after = block.get("space_after", None)
        alignment = block.get("alignment", None)  # left, center, right
        color = block.get("color", None)

        p = self.doc.add_paragraph()
        run = p.add_run(text)

        # Применяем форматирование из JSON
        if block.get("bold"):
            run.bold = True
        if block.get("italic"):
            run.italic = True
        if block.get("underline"):
            run.underline = True
        if font_name:
            run.font.name = font_name
        if font_size:
            from docx.shared import Pt

            run.font.size = Pt(font_size)
        if color:
            from docx.shared import RGBColor

            # Предполагаем, что цвет в формате "RRGGBB" или "RGB"
            if len(color) == 6:
                r = int(color[0:2], 16)
                g = int(color[2:4], 16)
                b = int(color[4:6], 16)
                run.font.color.rgb = RGBColor(r, g, b)

        # Применяем отступы и интервалы
        if (
            left_indent is not None
            or right_indent is not None
            or space_after is not None
        ):
            paragraph_format = p.paragraph_format
            if space_after is not None:
                from docx.shared import Pt

                paragraph_format.space_after = Pt(
                    space_after
                )  # Интервал после абзаца в пунктах

        # Проверяем, чтобы абзац не заходил на поля
        # (это также устанавливает отступы)
        self._ensure_paragraph_within_margins(
            p, left_indent, right_indent, alignment
        )

        # Выравнивание
        if alignment == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif alignment == "justify":
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # left - по умолчанию

    def _ensure_paragraph_within_margins(
        self, paragraph, left_indent=None, right_indent=None, alignment=None
    ):
        """
        Проверяет, чтобы абзац не заходил на поля документа.
        Корректирует отступы, если необходимо.
        """
        section = self.doc.sections[0]
        page_width = section.page_width
        left_margin = section.left_margin
        right_margin = section.right_margin

        # Если отступы не заданы, устанавливаем значения по умолчанию
        if left_indent is None:
            left_indent_inches = 0
        else:
            left_indent_inches = left_indent / 10.0  # Преобразуем в дюймы

        if right_indent is None:
            right_indent_inches = 0
        else:
            right_indent_inches = right_indent / 10.0  # Преобразуем в дюймы

        # Проверяем, чтобы суммарный отступ не превышал доступную ширину
        total_indents_inches = left_indent_inches + right_indent_inches
        available_width_for_content_inches = (
            page_width.inches - left_margin.inches - right_margin.inches
        )

        # Если отступы слишком велики для доступной ширины, корректируем их
        if total_indents_inches > available_width_for_content_inches:
            # Масштабируем отступы, чтобы они помещались на странице
            if total_indents_inches > 0:
                # Оставляем немного места для безопасности
                max_available_for_indents = (
                    available_width_for_content_inches * 0.8
                )
                scale_factor = max_available_for_indents / total_indents_inches
                if scale_factor < 1:
                    left_indent_inches = left_indent_inches * scale_factor
                    right_indent_inches = right_indent_inches * scale_factor
                    from docx.shared import Inches

                    paragraph.paragraph_format.left_indent = Inches(
                        left_indent_inches
                    )
                    paragraph.paragraph_format.right_indent = Inches(
                        right_indent_inches
                    )
        else:
            # Применяем отступы как обычно, если они в пределах допустимого
            from docx.shared import Inches

            paragraph.paragraph_format.left_indent = Inches(left_indent_inches)
            paragraph.paragraph_format.right_indent = Inches(
                right_indent_inches
            )

    def _list(self, block: dict):
        # Используем параметры из JSON, если они есть, иначе - стандартные
        ordered = block.get("ordered", False)
        font_name = block.get("font_name", None)
        font_size = block.get("font_size", None)
        left_indent = block.get("left_indent", None)
        right_indent = block.get("right_indent", None)
        space_after = block.get("space_after", None)
        alignment = block.get("alignment", "left")  # left, center, right
        color = block.get("color", None)

        style = "List Number" if ordered else "List Bullet"

        for item in block.get("items", []):
            cleaned_item = clean_html_tags(item)
            p = self.doc.add_paragraph(cleaned_item, style=style)

            # Применяем стили, если они указаны в JSON
            if p.runs:
                run = p.runs[0]
                if font_name:
                    run.font.name = font_name
                if font_size:
                    from docx.shared import Pt

                    run.font.size = Pt(font_size)
                if color:
                    from docx.shared import RGBColor

                    # Предполагаем, что цвет в формате "RRGGBB" или "RGB"
                    if len(color) == 6:
                        r = int(color[0:2], 16)
                        g = int(color[2:4], 16)
                        b = int(color[4:6], 16)
                        run.font.color.rgb = RGBColor(r, g, b)

            # Применяем отступы и интервалы
            if (
                left_indent is not None
                or right_indent is not None
                or space_after is not None
            ):
                paragraph_format = p.paragraph_format
                if left_indent is not None:
                    from docx.shared import Inches

                    paragraph_format.left_indent = Inches(
                        left_indent / 10.0
                    )  # Преобразуем в дюймы
                if right_indent is not None:
                    from docx.shared import Inches

                    paragraph_format.right_indent = Inches(
                        right_indent / 10.0
                    )  # Преобразуем в дюймы
                if space_after is not None:
                    from docx.shared import Pt

                    paragraph_format.space_after = Pt(
                        space_after
                    )  # Интервал после абзаца в пунктах

            # Выравнивание
            if alignment == "center":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif alignment == "right":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif alignment == "justify":
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def _table(self, block: dict):
        """
        Рендерим таблицу.
        JSON-схема:
        {
            "type": "table",
            "headers": ["Колонка1", "Колонка2"],
            "rows": [
                ["Значение1", "Значение2"]
            ]
        }
        """
        headers = block.get("headers", [])
        rows = block.get("rows", [])

        # Используем параметры из JSON, если они есть, иначе - стандартные
        table_params = block.get("params", {})
        header_font_name = table_params.get("header_font_name", None)
        header_font_size = table_params.get("header_font_size", None)
        header_bold = table_params.get(
            "header_bold", True
        )  # по умолчанию заголовки жирные
        header_italic = table_params.get("header_italic", False)
        header_color = table_params.get("header_color", None)
        body_font_name = table_params.get("body_font_name", None)
        body_font_size = table_params.get("body_font_size", None)
        body_bold = table_params.get("body_bold", False)
        body_italic = table_params.get("body_italic", False)
        body_color = table_params.get("body_color", None)
        table_style = table_params.get("table_style", "Table Grid")
        header_bg_color = table_params.get("header_bg_color", None)

        # Дополнительные параметры таблицы
        table_properties = block.get("table_properties", {})
        cell_properties_list = block.get("cell_properties", [])
        row_properties_list = block.get("row_properties", [])

        if not headers or not rows:
            return  # Пустая таблица — ничего не делаем

        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = table_style  # стиль таблицы из JSON

        # Устанавливаем ширину столбцов, если указана
        if "widths" in table_properties and table_properties["widths"]:
            widths = table_properties["widths"]
            for i, width_val in enumerate(widths):
                if i < len(table.columns):
                    try:
                        # Преобразуем TWIP в Inches (1 TWIP = 1/1440 дюйма)
                        from docx.shared import Inches

                        table.columns[i].width = Inches(width_val / 1440.0)
                    except (TypeError, ValueError):
                        # Игнорируем ошибки при установке ширины столбца
                        pass

        # Устанавливаем отступы ячеек, если указаны
        cell_margin = table_properties.get("cell_margin", {})
        if cell_margin:
            for row in table.rows:
                for cell in row.cells:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()

                    # Устанавливаем отступы
                    tc_mar = OxmlElement("w:tcMar")

                    for margin_type, margin_value in cell_margin.items():
                        margin_elem = OxmlElement(f"w:{margin_type}")
                        margin_elem.set(qn("w:w"), str(margin_value))
                        margin_elem.set(qn("w:type"), "dxa")
                        tc_mar.append(margin_elem)

                    tcPr.append(tc_mar)

        # Заголовки
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            cleaned_header = clean_html_tags(str(header))
            hdr_cells[i].text = cleaned_header

            # Применяем стили к заголовкам
            if (
                header_font_name
                or header_font_size
                or header_bold
                or header_italic
                or header_color
            ):
                run = (
                    hdr_cells[i].paragraphs[0].runs[0]
                    if hdr_cells[i].paragraphs
                    and hdr_cells[i].paragraphs[0].runs
                    else hdr_cells[i].paragraphs[0].add_run(cleaned_header)
                )
                if header_font_name:
                    run.font.name = header_font_name
                if header_font_size:
                    from docx.shared import Pt

                    run.font.size = Pt(header_font_size)
                if header_bold is not None:
                    run.font.bold = header_bold
                if header_italic is not None:
                    run.font.italic = header_italic
                if header_color:
                    from docx.shared import RGBColor

                    # Предполагаем, что цвет в формате "RRGGBB" или "RGB"
                    if len(header_color) == 6:
                        r = int(header_color[0:2], 16)
                        g = int(header_color[2:4], 16)
                        b = int(header_color[4:6], 16)
                        run.font.color.rgb = RGBColor(r, g, b)

        # Применяем фоновый цвет к заголовкам, если указан
        if header_bg_color:
            for i, header in enumerate(headers):
                # Получаем XML элемент ячейки
                tc = hdr_cells[i]._tc
                # Создаем элемент заливки
                tcFill = OxmlElement("w:shd")
                tcFill.set(qn("w:fill"), header_bg_color.replace("#", ""))
                # Добавляем заливку к ячейке
                tc.get_or_add_tcPr().append(tcFill)

        # Проверяем, чтобы таблица не выходила за пределы полей
        self._ensure_table_within_margins(table, table_properties)

        # Данные
        for row_idx, row in enumerate(rows):
            cells = table.add_row().cells
            # Применяем свойства строки, если они определены
            for row_prop in row_properties_list:
                if (
                    row_prop.get("row") == row_idx + 1
                ):  # +1 потому что 0-й ряд - заголовки
                    # Применяем заливку ко всей строке
                    if "bg_color" in row_prop:
                        bg_color = row_prop["bg_color"]
                        if bg_color.startswith("#"):
                            bg_color = bg_color[1:]  # Убираем #

                        for cell in cells:
                            tc = cell._tc
                            tcPr = tc.get_or_add_tcPr()

                            # Создаем элемент заливки
                            tc_fill = OxmlElement("w:shd")
                            tc_fill.set(qn("w:fill"), bg_color)
                            tcPr.append(tc_fill)

                    # Применяем цвет текста ко всей строке
                    if "text_color" in row_prop:
                        text_color = row_prop["text_color"]
                        if text_color.startswith("#"):
                            text_color = text_color[1:]  # Убираем #

                        from docx.shared import RGBColor

                        if len(text_color) == 6:
                            r = int(text_color[0:2], 16)
                            g = int(text_color[2:4], 16)
                            b = int(text_color[4:6], 16)
                            rgb_color = RGBColor(r, g, b)

                            # Применяем цвет ко всем ячейкам в строке
                            for cell in cells:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.color.rgb = rgb_color

            for col_idx, value in enumerate(row):
                cleaned_value = clean_html_tags(str(value))
                cells[col_idx].text = cleaned_value

                # Применяем стили к ячейкам данных
                if (
                    body_font_name
                    or body_font_size
                    or body_bold
                    or body_italic
                    or body_color
                ):
                    run = (
                        cells[col_idx].paragraphs[0].runs[0]
                        if cells[col_idx].paragraphs
                        and cells[col_idx].paragraphs[0].runs
                        else cells[col_idx]
                        .paragraphs[0]
                        .add_run(cleaned_value)
                    )
                    if body_font_name:
                        run.font.name = body_font_name
                    if body_font_size:
                        from docx.shared import Pt

                        run.font.size = Pt(body_font_size)
                    if body_bold is not None:
                        run.font.bold = body_bold
                    if body_italic is not None:
                        run.font.italic = body_italic
                    if body_color:
                        from docx.shared import RGBColor

                        # Предполагаем, что цвет в формате "RRGGBB" или "RGB"
                        if len(body_color) == 6:
                            r = int(body_color[0:2], 16)
                            g = int(body_color[2:4], 16)
                            b = int(body_color[4:6], 16)
                            run.font.color.rgb = RGBColor(r, g, b)

                # Применяем индивидуальные свойства ячейки
                for cell_prop in cell_properties_list:
                    if (
                        cell_prop.get("row") == row_idx + 1
                        and cell_prop.get("col") == col_idx
                    ):  # +1 потому что 0-й ряд - заголовки
                        self._apply_cell_properties(cells[col_idx], cell_prop)

    def _ensure_table_within_margins(self, table, table_properties):
        """
        Проверяет, чтобы таблица не выходила за пределы полей документа.
        Корректирует ширину столбцов, если необходимо.
        """
        section = self.doc.sections[0]
        page_width = section.page_width
        left_margin = section.left_margin
        right_margin = section.right_margin
        available_width_twips = (
            int((page_width - left_margin - right_margin).twips)
            if hasattr(page_width - left_margin - right_margin, "twips")
            else int(
                (
                    (
                        page_width.inches
                        - left_margin.inches
                        - right_margin.inches
                    )
                    * 1440
                )
            )
        )

        # Если заданы ширины столбцов в table_properties, проверяем их сумму
        if "widths" in table_properties and table_properties["widths"]:
            widths = table_properties["widths"]
            total_width_twips = sum(widths)  # Ширины в TWIP (1/20 пункта)

            # Если общая ширина таблицы больше доступной ширины страницы,
            # масштабируем
            if total_width_twips > available_width_twips:
                scale_factor = available_width_twips / total_width_twips
                for i, width_val in enumerate(widths):
                    if i < len(table.columns):
                        try:
                            scaled_width = int(width_val * scale_factor)
                            table.columns[i].width = scaled_width
                        except (TypeError, ValueError):
                            # Игнорируем ошибки при установке ширины столбца
                            pass
        else:
            # Если ширины не заданы, распределяем ширину равномерно
            num_cols = len(table.columns)
            if num_cols > 0:
                column_width_twips = int(available_width_twips / num_cols)
                for col in table.columns:
                    col.width = column_width_twips

    def _apply_cell_properties(self, cell, properties):
        """
        Применяет свойства к отдельной ячейке
        """
        try:
            # Применяем заливку ячейки
            if "bg_color" in properties:
                bg_color = properties["bg_color"]
                if bg_color.startswith("#"):
                    bg_color = bg_color[1:]  # Убираем #

                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()

                # Создаем элемент заливки
                tc_fill = OxmlElement("w:shd")
                tc_fill.set(qn("w:fill"), bg_color)
                tcPr.append(tc_fill)

            # Применяем перенос текста
            # В Word по умолчанию текст переносится,
            # поэтому если text_wrap=true, ничего не делаем
            # Если text_wrap=false, тогда отключаем перенос
            if "text_wrap" in properties and not properties["text_wrap"]:
                # Отключаем перенос текста
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()

                # Устанавливаем noWrap для отключения переноса
                no_wrap = OxmlElement("w:noWrap")
                tcPr.append(no_wrap)

            # Применяем выравнивание
            if "vertical_alignment" in properties:
                vertical_alignment = properties["vertical_alignment"]
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()

                # Устанавливаем вертикальное выравнивание
                v_align = OxmlElement("w:vAlign")
                if vertical_alignment == "top":
                    v_align.set(qn("w:val"), "top")
                elif vertical_alignment == "center":
                    v_align.set(qn("w:val"), "center")
                elif vertical_alignment == "bottom":
                    v_align.set(qn("w:val"), "bottom")
                else:
                    v_align.set(qn("w:val"), "center")  # по умолчанию

                tcPr.append(v_align)

            # Применяем цвет текста
            if "text_color" in properties:
                text_color = properties["text_color"]
                if text_color.startswith("#"):
                    text_color = text_color[1:]  # Убираем #

                # Применяем цвет ко всем параграфам в ячейке
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        from docx.shared import RGBColor

                        if len(text_color) == 6:
                            r = int(text_color[0:2], 16)
                            g = int(text_color[2:4], 16)
                            b = int(text_color[4:6], 16)
                            run.font.color.rgb = RGBColor(r, g, b)

            # Применяем горизонтальное выравнивание
            if "horizontal_alignment" in properties:
                horizontal_alignment = properties["horizontal_alignment"]
                p = (
                    cell.paragraphs[0]
                    if cell.paragraphs
                    else cell.add_paragraph()
                )

                if horizontal_alignment == "left":
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif horizontal_alignment == "center":
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif horizontal_alignment == "right":
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif horizontal_alignment == "justify":
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Применяем границы
            if "border" in properties:
                border_props = properties["border"]
                self._apply_border_to_cell(cell, border_props)
        except Exception:
            # Игнорируем ошибки при применении свойств ячейки
            pass

    def _apply_border_to_cell(self, cell, border_props):
        """
        Применяет границы к ячейке
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # Обрабатываем каждую границу
        for border_side in [
            "top",
            "bottom",
            "left",
            "right",
            "insideH",
            "insideV",
            "all",
        ]:
            if border_side in border_props:
                border_info = border_props[border_side]

                # Проверяем, что border_info - это словарь
                if not isinstance(border_info, dict):
                    continue

                # Создаем элемент границы
                border_elem = OxmlElement(f"w:{border_side}")

                # Устанавливаем стиль границы
                style = border_info.get("style", "single")
                border_elem.set(qn("w:val"), str(style))

                # Устанавливаем размер границы (в EIGHT_POINTS)
                size = border_info.get("size", 4)  # по умолчанию 4 = 0.5pt
                border_elem.set(qn("w:sz"), str(size))

                # Устанавливаем цвет границы
                color = border_info.get("color", "auto")
                border_elem.set(qn("w:color"), str(color))

                # Добавляем границу к свойствам ячейки
                tcPr.append(border_elem)

    def _math(self, block: dict):
        """
        Рендерим математическую формулу.
        JSON-схема:
        {
            "type": "math",
            "formula": "LaTeX formula",
            "caption": "optional caption"
        }
        """
        formula = block.get("formula", "")
        caption = block.get("caption", "")

        # Используем параметры из JSON, если они есть, иначе - стандартные
        font_name = block.get("font_name", None)
        font_size = block.get("font_size", None)
        # Используем font_size как fallback,
        # если конкретные параметры не указаны
        math_font_size = block.get("math_font_size", None) or font_size or 12
        caption_font_size = (
            block.get("caption_font_size", None) or font_size or 10
        )
        bold = block.get("bold", False)
        italic = block.get("italic", True)  # По умолчанию курсив для формул
        alignment = block.get("alignment", "left")  # left, center, right
        color = block.get("color", None)

        # Создаем изображение с формулой с помощью matplotlib
        # Это более надежный способ отображения сложных формул в Word
        try:
            # Настройка фигуры matplotlib с размерами,
            # пропорциональными длине формулы
            # Преобразуем размеры из миллиметров в дюймы (1 дюйм = 25.4 мм)
            base_width_mm = 60  # базовая ширина в мм
            width_factor = min(
                len(formula) / 10, 3
            )  # масштабируем в зависимости от длины формулы не более чем х3
            width_mm = base_width_mm * width_factor

            # Высота пропорциональна ширине, но с минимальным значением
            height_mm = max(width_mm * 0.2, 15)  # высота в мм, не менее 15 мм

            # Преобразуем миллиметры в дюймы для matplotlib
            width_inches = width_mm / 25.4
            height_inches = height_mm / 25.4

            fig, ax = plt.subplots(figsize=(width_inches, height_inches))
            ax.text(
                0.5,
                0.5,
                f"${formula}$",
                fontsize=16,
                ha="center",
                va="center",
                transform=ax.transAxes,
            )
            ax.axis("off")  # Скрыть оси

            # Сохраняем в байтовый поток
            img_buffer = io.BytesIO()
            plt.savefig(img_buffer, format="png", bbox_inches="tight", dpi=150)
            img_buffer.seek(0)

            # Добавляем изображение в документ
            paragraph = self.doc.add_paragraph()

            # Применяем выравнивание
            if alignment == "center":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif alignment == "right":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif alignment == "justify":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            run = paragraph.add_run()

            # Добавляем изображение в документ
            run.add_picture(img_buffer, width=Inches(6))

            # Если есть подпись, добавляем её
            if caption:
                caption_p = self.doc.add_paragraph()

                # Применяем выравнивание к подписи
                if alignment == "center":
                    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif alignment == "right":
                    caption_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif alignment == "justify":
                    caption_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                caption_run = caption_p.add_run(f"{caption}")

                # Применяем стили к подписи
                if font_name:
                    caption_run.font.name = font_name
                if caption_font_size:
                    from docx.shared import Pt

                    caption_run.font.size = Pt(caption_font_size)
                if bold:
                    caption_run.font.bold = bold
                if italic:
                    caption_run.font.italic = italic
                if color:
                    from docx.shared import RGBColor

                    # Предполагаем, что цвет в формате "RRGGBB" или "RGB"
                    if len(color) == 6:
                        r = int(color[0:2], 16)
                        g = int(color[2:4], 16)
                        b = int(color[4:6], 16)
                        caption_run.font.color.rgb = RGBColor(r, g, b)

        except Exception as e:
            # Если не удалось создать изображение формулы,
            # добавляем как обычный текст
            p = self.doc.add_paragraph()

            # Применяем выравнивание
            if alignment == "center":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif alignment == "right":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif alignment == "justify":
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            run = p.add_run(f"Формула: {formula}")

            # Применяем стили к формуле
            if font_name:
                run.font.name = font_name
            if math_font_size:
                from docx.shared import Pt

                run.font.size = Pt(math_font_size)
            if bold:
                run.font.bold = bold
            if italic:
                run.font.italic = italic
            if color:
                from docx.shared import RGBColor

                # Предполагаем, что цвет в формате "RRGGBB" или "RGB"
                if len(color) == 6:
                    r = int(color[0:2], 16)
                    g = int(color[2:4], 16)
                    b = int(color[4:6], 16)
                    run.font.color.rgb = RGBColor(r, g, b)

            print(f"Ошибка при рендеринге формулы: {e}")

            if caption:
                caption_p = self.doc.add_paragraph()

                # Применяем выравнивание к подписи
                if alignment == "center":
                    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif alignment == "right":
                    caption_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif alignment == "justify":
                    caption_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                caption_run = caption_p.add_run(f"({caption})")

                # Применяем стили к подписи
                if font_name:
                    caption_run.font.name = font_name
                if caption_font_size:
                    from docx.shared import Pt

                    caption_run.font.size = Pt(caption_font_size)
                if bold:
                    caption_run.font.bold = bold
                if italic:
                    caption_run.font.italic = italic
                if color:
                    from docx.shared import RGBColor

                    # Предполагаем, что цвет в формате "RRGGBB" или "RGB"
                    if len(color) == 6:
                        r = int(color[0:2], 16)
                        g = int(color[2:4], 16)
                        b = int(color[4:6], 16)
                        caption_run.font.color.rgb = RGBColor(r, g, b)

        finally:
            # Очищаем matplotlib
            plt.close()

    def _function_graph(self, block: dict):
        """
        Рендерим график математической функции.
        JSON-схема:
        {
            "type": "function_graph",
            "function": "mathematical function",
            "x_min": -10,
            "x_max": 10,
            "title": "Graph Title",
            "xlabel": "x",
            "ylabel": "y",
            "width": 6,
            "height": 4,
            "line_color": "blue",
            "line_width": 2,
            "show_grid": true,
            "caption": "optional caption",
            "alignment": "center"
        }
        """
        function_expr = block.get("function", "x")
        x_min = block.get("x_min", -10)
        x_max = block.get("x_max", 10)
        title = block.get("title", "")
        xlabel = block.get("xlabel", "x")
        ylabel = block.get("ylabel", "y")
        width = block.get("width", 6)
        height = block.get("height", 4)
        line_color = block.get("line_color", "blue")
        line_width = block.get("line_width", 2)
        show_grid = block.get("show_grid", True)
        caption = block.get("caption", "")
        alignment = block.get("alignment", "center")

        try:
            # Создаем массив x значений
            x = np.linspace(x_min, x_max, 400)

            # Подготавливаем выражение для вычисления
            # Проверяем, что выражение содержит только
            # допустимые символы для математических функций
            # Заменяем '^' на '**' для возведения в степень в Python
            # Убираем лишние пробелы
            safe_expr = function_expr.replace(" ", "")

            # Проверяем, что выражение не содержит LaTeX команды
            # или другие недопустимые элементы
            # Если выражение содержит LaTeX команды, это ошибка
            if (
                "\\equiv" in safe_expr
                or "\\pmod" in safe_expr
                or "\\int" in safe_expr
                or "\\sum" in safe_expr
                or "\\frac" in safe_expr
            ):
                raise ValueError(
                    f"Invalid function expression: '{function_expr}'."
                    f" This appears to be a LaTeX formula,"
                    f" not a mathematical expression."
                    f" Use 'math' block type for LaTeX formulas."
                )

            # Проверяем, что выражение не содержит знак присваивания
            # или другие недопустимые элементы
            if "=" in safe_expr:
                # Если в выражении есть '=', пытаемся извлечь правую часть
                parts = safe_expr.split("=", 1)
                if len(parts) > 1:
                    safe_expr = parts[
                        1
                    ].strip()  # Берем правую часть после знака равно
                else:
                    safe_expr = safe_expr.replace(
                        "=", ""
                    )  # Просто убираем знак равно если он не в уравнении

            # Заменяем '^' на '**' для возведения в степень в Python
            safe_expr = safe_expr.replace("^", "**")

            # Вычисляем y значения
            # Создаем безопасное окружение для eval
            namespace = {
                "x": x,
                "np": np,
                "sin": np.sin,
                "cos": np.cos,
                "tan": np.tan,
                "exp": np.exp,
                "log": np.log,
                "sqrt": np.sqrt,
                "abs": np.abs,
                "pi": np.pi,
                "e": np.e,
            }

            y = eval(safe_expr, {"__builtins__": {}}, namespace)

            # Создаем график
            fig, ax = plt.subplots(figsize=(width, height))
            ax.plot(x, y, color=line_color, linewidth=line_width)

            # Добавляем заголовок и подписи к осям
            if title:
                ax.set_title(title)
            ax.set_xlabel(xlabel)
            ax.set_ylabel(ylabel)

            # Показываем сетку, если нужно
            if show_grid:
                ax.grid(True)

            # Сохраняем в байтовый поток
            img_buffer = io.BytesIO()
            plt.savefig(img_buffer, format="png", bbox_inches="tight", dpi=150)
            img_buffer.seek(0)

            # Добавляем изображение в документ
            paragraph = self.doc.add_paragraph()

            # Применяем выравнивание
            if alignment == "center":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif alignment == "right":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif alignment == "justify":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            run = paragraph.add_run()

            # Добавляем изображение в документ
            run.add_picture(img_buffer, width=Inches(width))

            # Если есть подпись, добавляем её
            if caption:
                caption_p = self.doc.add_paragraph()

                # Применяем выравнивание к подписи
                if alignment == "center":
                    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif alignment == "right":
                    caption_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif alignment == "justify":
                    caption_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                caption_run = caption_p.add_run(f"{caption}")

                # Применяем стили к подписи
                # (аналогично реализации в _math методе)
                font_name = block.get("font_name")
                caption_font_size = block.get("caption_font_size", 10)
                bold = block.get("bold", False)
                italic = block.get("italic", False)
                color = block.get("color")

                if font_name:
                    caption_run.font.name = font_name
                if caption_font_size:
                    from docx.shared import Pt

                    caption_run.font.size = Pt(caption_font_size)
                if bold:
                    caption_run.font.bold = bold
                if italic:
                    caption_run.font.italic = italic
                if color:
                    from docx.shared import RGBColor

                    # Предполагаем, что цвет в формате "RRGGBB" или "RGB"
                    if len(color) == 6:
                        r = int(color[0:2], 16)
                        g = int(color[2:4], 16)
                        b = int(color[4:6], 16)
                        caption_run.font.color.rgb = RGBColor(r, g, b)

        except Exception as e:
            # Если не удалось создать график, добавляем сообщение об ошибке
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(
                f"Ошибка при построении графика функции '{function_expr}': {e}"
            )
            run.font.italic = True
            from docx.shared import RGBColor

            run.font.color.rgb = RGBColor(255, 0, 0)  # Красный цвет для ошибки

            print(f"Ошибка при рендеринге графика функции: {e}")

        finally:
            # Очищаем matplotlib
            plt.close()


def check_user_wants_word_format(user_message):
    """
    Check if user wants to receive description in Word format.

    Args:
        user_message (str): The message from user to analyze

    Returns:
        bool: True if user wants Word format, False otherwise
    """
    message = user_message.lower()
    return (
        "word" in message
        or "формат word" in message
        or "docx" in message
        or "в ворде" in message
        or "в формате документа" in message
        or "word документ" in message
        or "документ word" in message
        or "в microsoft word" in message
        or "в формате ворд" in message
        or "в формате word" in message
        or "в документе word" in message
        or "в word формате" in message
        or "в ворд формате" in message
        or "в вордовском формате" in message
        or "в формате ворда" in message
    )


async def send_docx_response(
    update,
    reply,
    image_url=None,
):
    """
    Отправляет DOCX-файл с ответом пользователю.
    Args:
        update: Объект обновления Telegram
        reply: Ответ от модели, который будет преобразован в DOCX
    """
    try:
        # Проверяем, что reply не пустой
        if not reply or reply.strip() == "":
            raise ValueError("Пустой ответ от модели")

        # Удаляем маркеры кода, если они есть
        cleaned_reply = reply.strip()
        if cleaned_reply.startswith("```json"):
            cleaned_reply = cleaned_reply[7:]  # Удаляем '```json'
        elif cleaned_reply.startswith("```"):
            cleaned_reply = cleaned_reply[3:]  # Удаляем '```'

        if cleaned_reply.endswith("```"):
            cleaned_reply = cleaned_reply[:-3]  # Удаляем закрывающий '```'

        cleaned_reply = cleaned_reply.strip()

        data = json.loads(cleaned_reply)
        doc_io = io.BytesIO()
        renderer = DocxRenderer()
        renderer.render(data, doc_io)
        doc_io.seek(0)

        await update.message.reply_document(
            document=InputFile(doc_io, filename="document.docx"),
            caption="Ваш ответ в формате Word",
        )
        if image_url:
            try:
                await update.message.reply_photo(
                    image_url,
                    caption="Вот что сгенерировано по Вашему запросу",
                )
            except TimedOut:
                await update.message.reply_text(
                    "⏰ Время ожидания отправки изображения истекло. "
                    "Изображение не было отправлено, но документ  создан."
                )
            except Exception as photo_error:
                if "timeout" in str(photo_error).lower():
                    await update.message.reply_text(
                        "⏰ Время ожидания отправки изображения истекло. "
                        "Изображение не было отправлено, но документ создан."
                    )
                else:
                    raise photo_error
    except json.JSONDecodeError as e:
        # Если ответ не является валидным JSON,
        # отправляем обычное сообщение
        from message_utils import send_long_message

        safe_reply = escape_markdown(reply, version=2)
        await send_long_message(update, safe_reply, parse_mode="MarkdownV2")
        print(f"Ошибка разбора JSON при создании DOCX: {e}")
    except ValueError as e:
        # Если возникла ошибка значения (например, пустой ответ)
        from message_utils import send_long_message

        safe_reply = escape_markdown(reply, version=2)
        await send_long_message(update, safe_reply, parse_mode="MarkdownV2")
        print(f"Ошибка значения при создании DOCX: {e}")
    except Exception as e:
        # Если не удалось создать или отправить
        # DOCX, отправляем обычное сообщение
        from message_utils import send_long_message

        safe_reply = escape_markdown(reply, version=2)
        await send_long_message(update, safe_reply, parse_mode="MarkdownV2")
        print(f"Ошибка при создании или отправке DOCX файла: {e}")
